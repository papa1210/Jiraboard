#!/usr/bin/env python3
"""
Convert .docx files to .md (Markdown)
Usage: python convert_docx_to_md.py <input.docx> [output.md]
If output is not specified, uses same name as input with .md extension.
"""

import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement
except ImportError:
    print("Error: python-docx not installed.")
    print("Install with: pip install python-docx")
    sys.exit(1)


def get_table_markdown(table):
    """Convert a docx table to markdown format."""
    md_lines = []
    
    # Header row
    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
    md_lines.append("| " + " | ".join(header_cells) + " |")
    md_lines.append("|" + "|".join(["---"] * len(header_cells)) + "|")
    
    # Body rows
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        md_lines.append("| " + " | ".join(cells) + " |")
    
    return "\n".join(md_lines)


def get_paragraph_style(paragraph):
    """Determine markdown style based on paragraph style."""
    style_name = paragraph.style.name.lower()
    
    if "heading 1" in style_name:
        return "#"
    elif "heading 2" in style_name:
        return "##"
    elif "heading 3" in style_name:
        return "###"
    elif "list" in style_name or paragraph.style.base_style and "list" in paragraph.style.base_style.name.lower():
        return "-"
    else:
        return None


def docx_to_markdown(docx_path):
    """Convert docx file to markdown content."""
    doc = Document(docx_path)
    md_lines = []
    
    for element in doc.element.body:
        # Handle paragraphs
        if element.tag.endswith('}p'):
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            
            if para is not None:
                text = para.text.strip()
                
                if not text:
                    md_lines.append("")
                    continue
                
                style = get_paragraph_style(para)
                
                if style == "-":
                    md_lines.append(f"- {text}")
                elif style:
                    md_lines.append(f"{style} {text}")
                else:
                    # Check for bold/italic formatting
                    if para.runs:
                        formatted_text = ""
                        for run in para.runs:
                            txt = run.text
                            if run.bold:
                                txt = f"**{txt}**"
                            if run.italic:
                                txt = f"*{txt}*"
                            formatted_text += txt
                        md_lines.append(formatted_text if formatted_text.strip() else "")
                    else:
                        md_lines.append(text)
        
        # Handle tables
        elif element.tag.endswith('}tbl'):
            for table in doc.tables:
                if table._element is element:
                    md_lines.append("")
                    md_lines.append(get_table_markdown(table))
                    md_lines.append("")
                    break
    
    return "\n".join(md_lines)


def main():
    if len(sys.argv) < 2:
        print("Usage: python convert_docx_to_md.py <input.docx> [output.md]")
        print("\nExample:")
        print("  python convert_docx_to_md.py 'Báo cáo cá nhân.docx'")
        print("  python convert_docx_to_md.py 'Báo cáo cá nhân.docx' 'Báo cáo cá nhân.md'")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    # Validate input file exists
    if not os.path.exists(input_file):
        print(f"Error: File '{input_file}' not found.")
        sys.exit(1)
    
    if not input_file.lower().endswith('.docx'):
        print("Error: Input file must be a .docx file.")
        sys.exit(1)
    
    # Determine output file
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        output_file = input_file.rsplit('.', 1)[0] + '.md'
    
    print(f"Converting: {input_file} → {output_file}")
    
    try:
        # Convert
        markdown_content = docx_to_markdown(input_file)
        
        # Write to output
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        print(f"✓ Success! Output saved to: {output_file}")
        
        # Show first 500 chars as preview
        preview = markdown_content[:500].replace('\n', '\n  ')
        print(f"\nPreview:\n  {preview}...")
        
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
